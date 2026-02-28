from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocumentFormatter:

    def __init__(self, lecture_title="Lecture Notes"):
        self.lecture_title = lecture_title

   
    # Public Method
    
    def generate_docx(self, structured_notes, metadata=None):
        """
        structured_notes format:
        [
            {
                "topic": str,
                "summary": str,
                "keywords": [str, str, ...]
            }
        ]
        """

        doc = Document()

        self._setup_styles(doc)
        self._add_title(doc)
        self._add_metadata(doc, metadata)

        for i, section in enumerate(structured_notes):
            self._add_section(doc, section)

            if i < len(structured_notes) - 1:
                doc.add_page_break()

        self._add_footer(doc)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


    # Styling Setup
    
    def _setup_styles(self, doc):

        styles = doc.styles

        # Body Style
        if "BodyStyle" not in styles:
            body_style = styles.add_style("BodyStyle", WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = "Times New Roman"
            body_style.font.size = Pt(12)

        # Keyword Style
        if "KeywordStyle" not in styles:
            keyword_style = styles.add_style("KeywordStyle", WD_STYLE_TYPE.PARAGRAPH)
            keyword_style.font.name = "Calibri"
            keyword_style.font.size = Pt(11)
            keyword_style.font.color.rgb = RGBColor(80, 80, 80)

  
    # Title
   
    def _add_title(self, doc):
        title = doc.add_heading(self.lecture_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph("")

    # Metadata Section
   
    def _add_metadata(self, doc, metadata):

        if not metadata:
            metadata = {}

        generated_on = datetime.now().strftime("%B %d, %Y")

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_text = f"Generated on: {generated_on}"
        if "duration" in metadata:
            meta_text += f" | Duration: {metadata['duration']}"
        if "model" in metadata:
            meta_text += f" | Model: {metadata['model']}"

        run = meta_para.add_run(meta_text)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph("")

    # Section Content
    
    def _add_section(self, doc, section):

        # Topic Heading
        heading = doc.add_heading(section["topic"], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = heading.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)

        # Summary Paragraph
        summary_para = doc.add_paragraph(section["summary"], style="BodyStyle")
        summary_para.paragraph_format.space_after = Pt(12)
        summary_para.paragraph_format.first_line_indent = Pt(18)

        # Keywords Label
        key_label = doc.add_paragraph("Key Concepts:", style="BodyStyle")
        key_label.runs[0].bold = True
        key_label.runs[0].font.color.rgb = RGBColor(204, 51, 0)

        # Keywords Inline
        keywords_text = ", ".join(section["keywords"])
        keywords_para = doc.add_paragraph(keywords_text, style="KeywordStyle")
        keywords_para.paragraph_format.space_after = Pt(14)

    # Footer + Page Number
    
    def _add_footer(self, doc):

        section = doc.sections[0]
        footer = section.footer

        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("Lumina AI | Page ")

        # Add page number field
        page_number = OxmlElement("w:fldSimple")
        page_number.set(qn("w:instr"), "PAGE")
        paragraph._p.append(page_number)

        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)